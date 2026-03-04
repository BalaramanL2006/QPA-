"""
Question Paper Difficulty Analyzer - Flask Backend
This application analyzes question papers and classifies questions based on difficulty level.
Supports multiple file formats: TXT, PDF, DOCX, and Images (JPG, PNG) with OCR.
Exclusively using Google Gemini API.
"""

from flask import Flask, render_template, request, jsonify, session, redirect, url_for  # type: ignore[import]
import sqlite3
import os
import time
from datetime import datetime
import json
from werkzeug.utils import secure_filename  # type: ignore[import]
from werkzeug.security import generate_password_hash, check_password_hash  # type: ignore[import]
import traceback
import re
from functools import wraps
from dotenv import load_dotenv  # type: ignore[import]
from google import genai  # type: ignore
from google.genai import types  # type: ignore

# 1. Gemini Configuration
load_dotenv()

# Initialize Gemini client using the NEW SDK
client = genai.Client(
    api_key=os.getenv("GEMINI_API_KEY")
)

# Direct imports for file handling libraries
import pdfplumber  # type: ignore[import]
from docx import Document  # type: ignore[import]
from PIL import Image  # type: ignore[import]
import pytesseract  # type: ignore[import]

app = Flask(__name__)

# Securely load secret key from environment
app.secret_key = os.getenv("FLASK_SECRET_KEY") or "fallback-unsecure-key"

# Configuration
DATABASE = 'database.db'
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'jpg', 'jpeg', 'png'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path):
    """Universal text extraction function."""
    try:
        file_ext = file_path.rsplit('.', 1)[1].lower()
        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext == 'pdf':
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        elif file_ext == 'docx':
            text = ""
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        elif file_ext in {'png', 'jpg', 'jpeg'}:
            image = Image.open(file_path)
            if image.mode != 'L':
                image = image.convert('L')
            return pytesseract.image_to_string(image)
        return ""
    except Exception as e:
        raise Exception(f"Extraction failed: {str(e)}")

def init_db():
    """Initialize the SQLite database."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS analysis_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            total_questions INTEGER,
            easy_count INTEGER,
            medium_count INTEGER,
            hard_count INTEGER,
            easy_percentage REAL,
            medium_percentage REAL,
            hard_percentage REAL,
            overall_difficulty TEXT,
            paper_text TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            school TEXT,
            mobile TEXT,
            password TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def get_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def analyze_with_ai(extracted_text):
    """
    Analyzes question paper text using the latest google-genai SDK.
    Forces strict JSON output using response_mime_type.
    """
    if not extracted_text:
        return {'error': 'No text provided for analysis.'}
        
    if not os.getenv("GEMINI_API_KEY"):
        return {'error': 'GEMINI_API_KEY is not configured.'}

    max_retries = 3
    base_delay = 2 # seconds

    for attempt in range(max_retries):
        try:
            # Construct Prompt (Strict JSON Format)
            prompt = f"""
You are an academic question paper analysis engine. Return a STRICT JSON response.
Required JSON format:
{{
  "total_questions": number,
  "total_marks": number,
  "question_types": {{ "mcq": number, "short_answer": number, "long_answer": number, "problem_solving": number }},
  "marks_distribution": {{ "1_mark": number, "2_mark": number, "5_mark": number, "10_mark": number, "others": number }},
  "difficulty_analysis": {{ "easy_percentage": number, "medium_percentage": number, "hard_percentage": number }},
  "bloom_taxonomy_distribution": {{ "remember": number, "understand": number, "apply": number, "analyze": number, "evaluate": number, "create": number }}
}}

Analyze this paper:
{extracted_text}
"""
            
            # Latest SDK Content Generation with Forced JSON
            response = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0,
                    max_output_tokens=4096
                )
            )
            
            if not response or not response.text:
                return {'error': 'Empty response from AI.'}
                
            raw_text = response.text.strip()
            
            # Safe JSON extraction using regex fallback
            try:
                return json.loads(raw_text)
            except json.JSONDecodeError:
                match = re.search(r'(\{.*\})', raw_text, re.DOTALL)
                if match:
                    try:
                        return json.loads(match.group(1))
                    except json.JSONDecodeError:
                        pass
                
                with open("ai_errors.log", "a", encoding="utf-8") as f:
                    f.write(f"\n[{datetime.now()}] JSON_DECODE_ERROR:\n{raw_text}\n{'-'*40}\n")
                return {'error': 'Invalid JSON response from AI.'}
                
        except Exception as e:
            err_msg = str(e)
            if "429" in err_msg and attempt < max_retries - 1:
                wait_time = base_delay * (2 ** attempt)
                print(f"Quota exceeded (429). Retrying in {wait_time}s... (Attempt {attempt + 1}/{max_retries})")
                time.sleep(wait_time)
                continue
                
            if "401" in err_msg: return {'error': 'Invalid API Key (401)'}
            if "404" in err_msg: return {'error': 'Model not found (404)'}
            if "429" in err_msg: return {'error': 'Quota exceeded (429). Please try again later.'}
            return {'error': f'API Error: {err_msg}'}
    
    return {'error': 'Maximum retries reached for AI analysis.'}

def save_to_history(entry):
    try:
        history_file = 'history.json'
        history_list = []
        if os.path.exists(history_file):
            with open(history_file, 'r', encoding='utf-8') as f:
                try:
                    history_list = json.load(f)
                except:
                    pass
        history_list.append(entry)
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history_list, f, indent=2, ensure_ascii=False)
        return True
    except:
        return False

def save_analysis(result, paper_text, filename):
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # Extract difficulty info from nested structure
        diff = result.get('difficulty_analysis', {})
        if not isinstance(diff, dict):
            diff = {}
            
        easy_p = diff.get('easy_percentage', 0)
        medium_p = diff.get('medium_percentage', 0)
        hard_p = diff.get('hard_percentage', 0)
        total_q = result.get('total_questions', 0)
        
        # Estimate counts (since NEW prompt only gives percentages)
        easy_c = int((easy_p / 100) * total_q) if total_q else 0
        medium_c = int((medium_p / 100) * total_q) if total_q else 0
        hard_c = int((hard_p / 100) * total_q) if total_q else 0

        cursor.execute('''
            INSERT INTO analysis_history 
            (filename, total_questions, easy_count, medium_count, hard_count, 
             easy_percentage, medium_percentage, hard_percentage, 
             overall_difficulty, paper_text)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            filename or 'Text Input',
            total_q,
            easy_c,
            medium_c,
            hard_c,
            easy_p,
            medium_p,
            hard_p,
            "Analyzed", # Generic label since NEW prompt removed overall_difficulty
            paper_text
        ))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"DB Error: {e}")

# Auth Decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Routes
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name, email, school, mobile, password = request.form.get('name'), request.form.get('email'), request.form.get('school'), request.form.get('mobile'), request.form.get('password')
        hashed_password = generate_password_hash(password)
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute('INSERT INTO users (name, email, school, mobile, password) VALUES (?, ?, ?, ?, ?)', (name, email, school, mobile, hashed_password))
            conn.commit()
            conn.close()
            return redirect(url_for('login'))
        except:
            return render_template('register.html', error="Email already exists.")
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email, password = request.form.get('email'), request.form.get('password')
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()
        conn.close()
        if user and check_password_hash(user['password'], password):
            session['user_id'], session['user_name'] = user['id'], user['name']
            return redirect(url_for('staff_details'))
    return render_template('login.html')

@app.route('/staff-details', methods=['GET', 'POST'])
@login_required
def staff_details():
    if request.method == 'POST':
        session['staff_name'], session['subject'], session['subject_code'], session['syllabus'] = request.form.get('staff_name'), request.form.get('subject'), request.form.get('subject_code'), request.form.get('syllabus')
        return redirect(url_for('home'))
    return render_template('staff_details.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/')
@login_required
def home():
    if 'staff_name' not in session: return redirect(url_for('staff_details'))
    return render_template('index.html', staff_name=session.get('staff_name'), subject=session.get('subject'), subject_code=session.get('subject_code'), syllabus=session.get('syllabus', ''))

@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    try:
        paper_text, filename = "", ""
        if 'file' in request.files:
            file = request.files['file']
            if file.filename != '' and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(path)
                paper_text = extract_text(path)
                os.remove(path)
                # Call AI with full text
                analysis_result = analyze_with_ai(paper_text)
        elif request.is_json:
            data = request.get_json()
            # If standard JSON input, join questions into text
            questions = data.get('questions', [])
            filename = data.get('filename', 'API Input')
            paper_text = "\n".join(questions)
            analysis_result = analyze_with_ai(paper_text)
        else:
            return jsonify({'error': 'Invalid request'}), 400

        # Print the full AI error in terminal before returning 500
        if 'error' in analysis_result:
            print("AI ERROR:", analysis_result)
            return jsonify(analysis_result), 500

        # Mapping for backward compatibility with history system
        diff = analysis_result.get('difficulty_analysis', {})
        if not isinstance(diff, dict):
            diff = {}
        
        # Build history entry
        history_entry = {
            'filename': filename or 'Analysis',
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'level': "Analyzed",
            'total_questions': analysis_result.get('total_questions', 0),
            'easy': diff.get('easy_percentage', 0),
            'medium': diff.get('medium_percentage', 0),
            'hard': diff.get('hard_percentage', 0)
        }
        
        save_to_history(history_entry)
        save_analysis(analysis_result, paper_text, filename)
        
        return jsonify(analysis_result), 200
        
    except Exception as e:
        import traceback
        print("UNEXPECTED ERROR:")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/history')
@login_required
def history():
    history_data = []
    if os.path.exists('history.json'):
        with open('history.json', 'r', encoding='utf-8') as f:
            try: history_data = json.load(f)[::-1]
            except: pass
    return render_template('history.html', history=history_data)

@app.route('/clear_all_history', methods=['POST'])
def clear_all_history():
    with open('history.json', 'w', encoding='utf-8') as f: json.dump([], f)
    return redirect('/history')

if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='127.0.0.1', port=5000)
